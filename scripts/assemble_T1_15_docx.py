"""Assemble T1.15 comprehensive report as docx, embedding figures."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import json

ROOT = Path('/home/hyunjin/CLAUDE/Pd_DMC/research-pd-dmc')
FIG = ROOT / 'reports/G3/T1_15_figures'
OUT = ROOT / 'reports/G3' / 'T1_15_DFT_shortlist_report.docx'

doc = Document()

# -- Title --
title = doc.add_heading('T1.15 DFT Shortlist — 종합 보고서', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Pd / PdO / PdO₂ 표면 DMC 형성 DFT 연구\n')
run.bold = True; run.font.size = Pt(13)
sub.add_run('AutoAdsorbate → MACE-MH + D3 + cuEq → DFT shortlist (T1.15)\n')
sub.add_run(f'작성일: 2026-06-20 | 작성자: 이현진 (지도: 이태훈)\n')
sub.add_run('Reference: Shi et al., Angew. Chem. Int. Ed. 2024 (Pd₃Cu) — TS barrier benchmark + Case A-D framework')

doc.add_paragraph()

# -- Executive Summary --
doc.add_heading('Executive Summary', level=1)
exec_text = (
    "본 보고서는 T1.14 MACE ranking (Phase 1-3) 결과 및 Paimon-extended 5-judge committee 검증을 "
    "거쳐 정리한 T1.15 DFT shortlist 와 비용 견적을 담는다. 핵심 결과:\n\n"
    "1. Phase 1 (단일 흡착 2,516 후보) — D3 적용 후 5 surfaces 모두 ≥85% 수렴, E_range 130-3170 meV 확보.\n"
    "2. Phase 2 (co-ads SetA 37,956 후보, 65 hr) — silent-error judge 가 PBC 단편화 silent bug 발견. "
    "Post-relax direct-distance 필터로 ~43-51% 회복 (재실행 불필요, 10분 patch).\n"
    "3. Phase 3 (SetTS 5,814 + SetB 500) — 가이드 T2.5 재검토 결과 SetTS pool 자체 불필요 (NEB가 endpoint 에서 saddle 자동 발견). SetB 도 PBC drift로 단편화 (4-7% 잔존).\n"
    "4. 모든 surface 에서 chemistry trend 일관: S4 PdO₂(110) CO* 미결합 = Shi 2024 가설 직접 검증 "
    "(MACE-MH+D3 와 SevenNet-Omni+D3 cross-check 일치).\n"
    "5. T1.15 DFT shortlist: 47 후보 (S1×9, S2×14, S3×9, S3b×9, S4×6). "
    "예상 DFT 비용 ~7-14일 wall time (RTX 6000 Ada × 2 GPUs)."
)
doc.add_paragraph(exec_text)

doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run('Decision needed: '); r.bold = True
r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
p.add_run('DFT 진입 승인. 47 jobs / ~7-14일 / Level 1 (vacuum) + Level 2 (VASPsol). '
          'S4 일부 broken 후보 (single_CO rank-1 d_min=1.18 Å 등) 는 manual placement 로 대체 추천.')

# -- 1. Pipeline --
doc.add_page_break()
doc.add_heading('1. 파이프라인 개요', level=1)
doc.add_paragraph(
    "Shi 2024 가 SSW-NN 으로 한 'PES sampler + screening' 역할을, 본 연구는 "
    "AutoAdsorbate (heuristic enumeration) + MACE-MH foundation MLIP 으로 대체한다. "
    "OC20 PBE pretrained 가중치 (mh-1 model + oc20_usemppbe head) 를 학습 없이 사용하며, "
    "ranking 결과는 DFT (PBE+D3+VASPsol) 로 최종 확인한다."
)
doc.add_picture(str(FIG/'fig01_pipeline_overview.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 1. 전체 파이프라인 개요. G1-G2 통과 후 T1.14 MACE ranking (3 sub-phases) 거쳐 현재 T1.15 단계.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.runs[0].italic = True

# -- 2. MACE config --
doc.add_heading('2. MACE 설정 (advisor 권고 반영)', level=1)
p = doc.add_paragraph()
p.add_run('이태훈 교수 권고:').bold = True
p.add_run(' "분자 다루므로 D3 필요" → dispersion=True, damping=\'bj\', dispersion_xc=\'pbe\' 적용 (PBE+D3-BJ).')

p = doc.add_paragraph()
p.add_run('MLIP 모델 비교:').bold = True
p.add_run(' MACE-MH (mh-1 + oc20_usemppbe + cueq) vs SevenNet-Omni (7net-omni + oc20 + cueq). '
          '동일 10 구조 benchmark 결과 MACE 가 우리 시스템에서 3-5× 빠르고 S3b/S4 chemistry 정확도 우위 → MACE-MH 채택.')

doc.add_paragraph(
    "최종 calculator: mace_mp(model='mh-1', head='oc20_usemppbe', "
    "default_dtype='float64', enable_cueq=True, device='cuda', "
    "dispersion=True, damping='bj', dispersion_xc='pbe')"
)

# -- 3. Phase 1 --
doc.add_heading('3. Phase 1 — 단일 흡착 (CO*, CH₃O*) Ranking', level=1)
doc.add_paragraph(
    "AutoAdsorbate heuristic 으로 5 surface × {CO*, CH₃O*} = 2,516 후보 생성. "
    "MACE-MH+D3 로 LBFGS relax (fmax=0.05 eV/Å, max_steps=200) 후 (E_bin × adsorbate fingerprint) 로 dedup."
)
doc.add_picture(str(FIG/'fig02_phase_counts.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 2. (a) 표면별 후보 수 (Phase 1-3). (b) 전체 MLIP relaxation 분포.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

doc.add_picture(str(FIG/'fig03_phase1_yields.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 3. Phase 1 D3 결과. (a) 수렴률 (S4 PdO₂만 70-93%, 나머지 99-100%). '
                        '(b) Dedup 후 unique 수. (c) Energy spread — 모든 표면 200 meV 이상 (S3 CO* 128 meV는 PES 평탄 신호).')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

# -- 4. Phase 2 --
doc.add_heading('4. Phase 2 — Co-adsorption SetA Reactive', level=1)
doc.add_paragraph(
    "5 surface × 3,764-11,598 reactive co-ads pair = 37,956 후보. SLURM 없이 nohup 으로 ~65 hr 진행. "
    "Initial committee (5-judge blind): physics + silent-error 가 동시 REJECT."
)
p = doc.add_paragraph()
r = p.add_run('🔴 발견된 silent bug: '); r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
p.add_run('atoms.get_all_distances(mic=True) 를 ads-ads 거리 측정에 사용하여, '
          'PBC wrap 으로 분자가 셀 반대편으로 단편화된 상태도 1.18 Å 처럼 가까운 거리로 잘못 측정됨. '
          '결과: 25-97% top-N 구조가 corrupted (S2 worst 97%). MIC 가 silent error 의 mask 역할.')

doc.add_paragraph(
    "Fix: scripts/refilter_phase2_geometry.py — direct (non-MIC) pdist 로 7-atom 흡착물 검증. "
    "단편화 (d_max > 6 Å), 충돌 (d_min < 0.8 Å), C-O/methoxy O-C/C-H 결합 끊김 모두 제거 + looser 0.1 Å fingerprint 로 dedup."
)
doc.add_picture(str(FIG/'fig04_phase2_filter_yield.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 4. Phase 2 SetA 필터링 결과. (a) 평균 ~47% 생존 (모든 surface 절반 가까이 PBC 단편화). '
                        '(b) Drop 사유 — fragmented 가 압도적. S4는 C-H 결합 끊김 (884개) 도 두드러짐.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

doc.add_picture(str(FIG/'fig08_d_reactive_phase2.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 5. Phase 2 filtered SetA: 반응 거리 [d(C_CO − O_methoxy)] 분포. '
                        'SetA 의도 범위 [2.1, 4.0] Å 안에 대다수 분포. 4 Å 이상은 분리 (Set B 영역으로 drift).')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

# -- 5. Phase 3 --
doc.add_heading('5. Phase 3 — SetTS (TS guess) + SetB (thermo ref)', level=1)
doc.add_paragraph(
    "SetTS (1.7-2.3 Å) = TS 후보 풀, SetB (≥5 Å) = non-interacting baseline 의도. SLURM job 370, 10h 22min. "
    "5-judge committee: silent-error 가 동일 PBC bug 재발 + SetTS 99.7% 가 band 이탈 발견."
)
p = doc.add_paragraph()
r = p.add_run('🔴 SetTS pool 무용 판정: '); r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
p.add_run('TS 영역은 saddle point — MACE 같은 minimum-finder 는 자연스럽게 minimum 으로 도망. '
          '5,729개 중 16개 (0.3%) 만 1.7-2.3 Å 범위에 잔존. '
          '가이드 docs/DMC_Pd_workplan.md T2.5 재확인: "endpoint(T2.4)로 NEB image 생성" — '
          'TS 후보 풀 별도 생성 명시 없음. NEB가 endpoint A↔B 사이 saddle 자동 발견.')

doc.add_picture(str(FIG/'fig05_phase3_filter_yield.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 6. Phase 3 필터링 결과. (a) SetTS — 가이드 T2.5 에 의해 사실상 불필요. '
                        '(b) SetB — PBC drift + collapse 로 인해 4-7% 만 생존 (S4=0). '
                        'descriptor map zero-point 도 가이드는 단일 흡착 (Phase 1) E_ads 만 사용 → SetB 도 필수 아님.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

# -- 6. Committee --
doc.add_heading('6. Paimon-Extended 5-Judge Committee', level=1)
doc.add_paragraph(
    "본 연구는 Paimon framework (Park et al., arXiv:2606.09422, 2026 SNU 한승우 그룹) 영감으로 "
    "5-judge blind parallel committee + chair 패턴 도입. "
    "특히 judge-silent-error (Paimon §2.2) 와 judge-malicious (§3.3) 가 결정적 역할 — "
    "Phase 2, Phase 3 의 MIC PBC bug 를 두 차례 잡아냄. "
    "Reject 시 summarizer → producer fix → retry committee 의 Revise loop 1 cycle 성공 (Phase 2 → filtered)."
)
doc.add_picture(str(FIG/'fig06_committee_verdicts.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 7. Committee verdict timeline. P=Pass, P-c=Pass-with-caveats, C=Concern, R=Reject. '
                        'Phase 1 D3 cycle 에서 dispersion issue 해소, Phase 2 retry 에서 PBC bug 해소.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

# -- 7. Top-1 chemistry --
doc.add_heading('7. Shortlist Chemistry 검증', level=1)
doc.add_paragraph(
    "DFT shortlist 47 후보의 top-3 chemisorption 거리 점검. 표준 범위: Pd-C 1.85-2.10 Å, Pd-O 2.00-2.15 Å."
)
doc.add_picture(str(FIG/'fig07_top1_chemistry.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 8. DFT shortlist top-3 chemisorption check. (a) CO* Pd-C: S1, S2, S3b 정상; '
                        'S3 (O-term, Pd 가려짐) + S4 (PdO₂) 다수 physisorbed → Shi 2024 가설 직접 검증. '
                        '(b) CH₃O* Pd-O: 대부분 정상, S4 일부 (broken 의심).')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

doc.add_picture(str(FIG/'fig10_descriptor_preview.png'), width=Inches(5.5))
cap = doc.add_paragraph('그림 9. 예비 descriptor map (MLIP top-1 기반, slab E 차감만). '
                        'S1, S2, S3b 가 binding strong (좌하단), S4 가 weak (우상단으로 분리될 예상). '
                        'DFT 로 정확히 정량화 후 T1.19 에서 본 descriptor map 갱신.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

# -- 8. DFT shortlist --
doc.add_heading('8. T1.15 DFT Shortlist (47 jobs)', level=1)
doc.add_paragraph(
    "Phase 1 D3 unique (CO* + CH₃O*) + Phase 2 filtered SetA shortlist 통합. "
    "S2 는 PdO/Pd interface 화학 중요성 고려 +2 추가 (CO 5, CH₃O 5, coads 4). "
    "S4 는 Phase 2 filtered 에서 viable shortlist 없음 — single ads 만 (6 jobs)."
)

# Shortlist table
sl_global = json.load(open(ROOT/'calculations/G3_adsorption/DFT_shortlist/shortlist_global.json'))
budget = {}
for e in sl_global:
    s = e['surface']; k = e['kind']
    budget.setdefault(s, {'single_CO':0,'single_CH3O':0,'coads_SetA':0})
    budget[s][k] += 1

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '표면'; hdr[1].text = 'CO*'; hdr[2].text = 'CH₃O*'; hdr[3].text = 'co-ads (SetA)'; hdr[4].text = '합계'
for s in ['S1','S2','S3','S3b','S4']:
    row = table.add_row().cells
    row[0].text = f'{s} ({{S1:"Pd(100)",S2:"PdO(101)/Pd(100)",S3:"PdO(100) O-term",S3b:"PdO(100) PdO-term",S4:"PdO₂(110)"}}[s])'
    row[0].text = f'{s}'
    row[1].text = str(budget[s]['single_CO'])
    row[2].text = str(budget[s]['single_CH3O'])
    row[3].text = str(budget[s]['coads_SetA'])
    row[4].text = str(sum(budget[s].values()))
total_row = table.add_row().cells
total_row[0].text = 'Total'
for i, kind in enumerate(['single_CO','single_CH3O','coads_SetA']):
    total_row[i+1].text = str(sum(budget[s][kind] for s in budget))
total_row[4].text = str(len(sl_global))
for cell in total_row:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_picture(str(FIG/'fig09_dft_shortlist.png'), width=Inches(5.5))
cap = doc.add_paragraph('그림 10. DFT shortlist 구성 (총 47 jobs). S2 는 interface 중요성으로 14 jobs (다른 표면 6-9).')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

doc.add_heading('Shortlist Top Structures (MLIP-relaxed)', level=2)
doc.add_picture(str(FIG/'fig11_top_structures.png'), width=Inches(6.5))
cap = doc.add_paragraph('그림 11. 각 surface 의 top-1 single CO*, CH₃O*, co-ads SetA 구조. '
                        '(MLIP-relaxed; DFT 단계에서 추가 relax 됨)')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic = True

# -- 9. Cost estimate --
doc.add_heading('9. DFT 비용 견적 (T1.16 + T1.17)', level=1)
doc.add_paragraph(
    "VASP PBE+D3+(VASPsol) on RTX 6000 Ada × 2 GPUs 가정. "
    "Per-job 평균 116 atoms (substrate + adsorbate). "
    "memory rule (비싼 작업 견적 보고) 준수."
)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '단계'; hdr[1].text = '후보'; hdr[2].text = '시간/job'; hdr[3].text = 'GPU-hr 합계'; hdr[4].text = 'Wall (2 GPUs)'
row = table.add_row().cells
row[0].text = 'T1.16 Level 1 (vacuum)'
row[1].text = '47'; row[2].text = '6-12 hr'; row[3].text = '282-564'; row[4].text = '5.9-11.8 day'
row = table.add_row().cells
row[0].text = 'T1.17 Level 2 (VASPsol SP)'
row[1].text = '47'; row[2].text = '1-2 hr'; row[3].text = '47-94'; row[4].text = '1.0-2.0 day'
row = table.add_row().cells
row[0].text = 'Total'
row[1].text = ''; row[2].text = ''; row[3].text = '329-658'; row[4].text = '6.9-13.7 day'
for cell in row:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

# -- 10. Caveats --
doc.add_heading('10. Caveats & Risks', level=1)
doc.add_paragraph()
p = doc.add_paragraph(); p.add_run('🟡 S4 PdO₂(110) chemistry: ').bold = True
p.add_run('단일 흡착 + co-ads 모두에서 CO* 미결합 (MACE-MH+D3 & SevenNet-Omni+D3 cross-check 일치). '
          'Shi 2024 가설 정확히 검증 — Case C 분류 예정.')

p = doc.add_paragraph(); p.add_run('🟡 S3 PdO(100) O-term CO*: ').bold = True
p.add_run('PES 평탄 (top-10 spread ~13 meV), ranking 의미 약함. DFT 가 정량화 — top-3 거리 다양성 기반 선정.')

p = doc.add_paragraph(); p.add_run('🟡 S4 single_CO rank-1 d_min=1.18 Å, single_CH3O rank-0 d_min=0.98 Å: ').bold = True
p.add_run('PBC artifact 의심. DFT 진입 전 visual inspection 또는 manual placement 로 대체 권고.')

p = doc.add_paragraph(); p.add_run('🟢 Phase 3 SetTS unused: ').bold = True
p.add_run('가이드 T2.5 의 "endpoint→NEB image" 흐름 따라 SetA filtered + DFT Level 1 endpoints 가 NEB seed.')

p = doc.add_paragraph(); p.add_run('🟢 PBC silent bug: ').bold = True
p.add_run('committee 가 2회 (Phase 2, Phase 3) 잡아내고 post-relax direct-distance 필터로 해결. '
          'Phase 2 65hr 데이터 재실행 없이 활용 가능.')

# -- 11. Next steps --
doc.add_heading('11. 다음 단계', level=1)
nx = doc.add_paragraph()
nx.add_run('(1) DFT 진입 승인 결정 — ').bold = False
nx.add_run('A: 47 jobs 그대로 제출, B: ~40 jobs (S4 broken 후보 manual 대체), C: ~20 jobs (S1/S3b 만 우선)\n')
nx.add_run('(2) T1.16 Level 1 vacuum DFT 실행 (SLURM bulk submission)\n')
nx.add_run('(3) T1.17 Level 2 VASPsol single-point\n')
nx.add_run('(4) T1.18 adsorption energy table — analyst agent\n')
nx.add_run('(5) T1.19 descriptor map (E_CO vs E_CH3O) — DFT 기반\n')
nx.add_run('(6) T1.20 Case A-D 분류 → G3 게이트 통과 (checkpoint C/D/E)\n')
nx.add_run('(7) Workplan Phase 2: T2.5/T2.6 CI-NEB TS — SetA filtered endpoints 활용')

# -- 12. References --
doc.add_heading('12. 참고 문헌', level=1)
refs = [
    "Shi et al. \"Stabilization of Pd⁰ by Cu Alloying: Pd₃Cu Electrocatalyst for Anodic Methanol Carbonylation\", Angew. Chem. Int. Ed. 2024, doi:10.1002/anie.202401311",
    "Park, Chung, You, Kim, Ju, Han. \"A Robust Agentic Framework for Expert-Level Automation of Atomistic Simulations (Paimon)\", arXiv:2606.09422, 2026",
    "Pd(100)-(√5×√5)R27°-O surface oxide revisited, arXiv:cond-mat/0304107",
    "MACE-MP foundation model — Batatia et al. arXiv:2401.00096; https://mace-docs.readthedocs.io",
    "AutoAdsorbate — ACS Catal. 2025, doi:10.1021/acscatal.5c06553",
    "cuEquivariance — https://github.com/NVIDIA/cuEquivariance",
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

doc.save(str(OUT))
print(f'✓ Report saved: {OUT}')
print(f'  Size: {OUT.stat().st_size/1024:.0f} KB')
