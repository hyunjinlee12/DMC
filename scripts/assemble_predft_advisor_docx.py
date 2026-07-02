"""Assemble pre-DFT advisor report docx.

Style: Arial body, clean figure layout (no figure titles), all captions Arial.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path('/home/hyunjin/CLAUDE/Pd_DMC/research-pd-dmc')
FIG = ROOT / 'reports/predft_advisor_figures'
OUT = ROOT / 'reports/PreDFT_Final_Report_Arial.docx'

doc = Document()

# Set body style to Arial
def set_arial_style(doc):
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

set_arial_style(doc)


def add_para(text, bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(10); r.font.name = 'Arial'


def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'


# =========== TITLE ===========
add_h('Pd / PdO / PdO₂ DMC Formation DFT Project', 0)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('Pre-DFT Final Report — DFT Submission Ready (T1.16)',
         bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('작성일: 2026-06-24 | 작성자: 이현진 (지도: 이태훈 교수, DGIST DDMM)',
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Anchor: Shi et al., Angew. Chem. Int. Ed. 2024 (Pd₃Cu anodic methanol carbonylation)',
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Framework: Paimon-extended 5-judge committee (Park et al., arXiv:2606.09422, 2026)',
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('GitHub: hyunjinlee12/DMC + DDMM-dgist/Pd-DMC (mirror, Pattern A 브랜치 관리)',
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# =========== EXECUTIVE SUMMARY ===========
add_h('Executive Summary', 1)
add_para(
    "본 보고서는 G1 → G2 → G3 (T1.10-T1.15) 단계 완료 + T1.16 DFT 진입 직전 상태를 "
    "종합하여 advisor 검토를 요청한다. Paimon-extended 5-judge committee 가 11개 "
    "verification cycle 을 수행했으며, 두 차례의 silent bug 를 잡고 fix 함으로써 "
    "65 시간짜리 MLIP 데이터를 재실행 없이 재활용. 가이드 docs/DMC_Pd_workplan.md §P1-C "
    "의 ranking 우선순위 5개를 strict 적용한 v2 shortlist (48 candidates) 가 DFT 입력으로 준비됨."
)
add_para('진행 단계 요약:', bold=True, size=12)
add_para(
    "  • G1 bulk (Pd, PdO, PdO₂)              : ✓ 통과 (Pass-with-caveats)\n"
    "  • G2 slab (5 surfaces)                  : ✓ 통과 (Pass-with-caveats + literature 일치)\n"
    "  • G3 T1.10-T1.14 MLIP ranking            : ✓ 완료 (MACE-MH+D3+cueq, ~48k 후보)\n"
    "  • G3 T1.15 DFT shortlist (v2 strict)    : ✓ 완료 (48 candidates)\n"
    "  • G3 T1.16 Level 1 vacuum DFT            : ⏳ 진입 대기 (user 승인 요청)\n"
    "  • G3 T1.17-T1.20 (VASPsol → Case A-D)    : 예정\n"
    "  • G4 workplan Phase 2 (NEB)              : 예정"
)
add_para('주요 chemistry 발견 (MLIP 단계, DFT 로 확정 예정):', bold=True, size=12)
add_para(
    "  ① Shi 2024 가설 직접 검증: Pd⁰→Pd⁴⁺ 산화 시 CO* 결합 약화 (S1 Pd-C 1.97 Å → S4 4.05 Å)\n"
    "  ② S3 PdO(100) O-term 에서 CH₃OCO* product 형성 신호 (d_react 1.34 Å) — Shi 미언급 신규\n"
    "  ③ S2/S4 에서 CO + lattice O → CO₂-like side-path 자발 sampling\n"
    "  ④ Cross-MLIP validation: MACE-MH+D3 vs SevenNet-Omni+D3 일치"
)
add_para('Decision request:', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
add_para(
    "T1.16 Level 1 vacuum DFT 진입 승인 (47-48 jobs, ~7-14 day wall time on 2 GPUs)."
)

# =========== 1. PIPELINE & METHOD ===========
doc.add_page_break()
add_h('1. Pipeline & Method', 1)
add_para(
    "본 연구는 Shi 2024 의 SSW-NN + DFT 접근을 AutoAdsorbate (heuristic enumeration) + "
    "MACE-MH foundation MLIP (mh-1 + oc20_usemppbe head) + DFT 의 3단 파이프라인으로 대체. "
    "Foundation MLIP 는 OC20 PBE 사전학습 가중치를 그대로 사용 (system-specific 학습 불필요). "
    "각 단계의 출력은 Paimon-extended 5-judge committee 가 blind parallel review 로 검증."
)
add_para('MACE calculator 설정 (advisor 권고 반영):', bold=True, size=12)
add_para(
    "  mace_mp(model='mh-1',                  # MACE-MH multihead foundation\n"
    "          head='oc20_usemppbe',          # Open Catalyst 20 PBE\n"
    "          default_dtype='float64',        # 정확도 (이용혁 교수 조언)\n"
    "          enable_cueq=True,               # cuEquivariance 가속 (~2× 빠름)\n"
    "          dispersion=True,                # D3 dispersion (advisor 권고)\n"
    "          damping='bj', dispersion_xc='pbe')"
)

# =========== 2. SURFACE MODELS ===========
add_h('2. Surface Models (5 slabs)', 1)
add_para(
    "Pd⁰ → Pd⁴⁺ 산화 진행 축을 따라 5개 slab. 각 표면은 published DFT literature 와 sub-percent "
    "lattice 일치 (Reuter PRL 2007, Lundgren cond-mat/0304107, Rogal-Reuter PRB 2004). "
    "S4 PdO₂(110) 만 직접 DFT benchmark 없음 (rutile TiO₂/RuO₂ 유추) — Shi 2024 의 Pd⁴⁺ "
    "representation 으로 해석."
)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['ID', '표면 모델', '산화 상태', 'Atoms', 'E (E₀, eV)']):
    hdr[i].text = h
for r in [
    ('S1', 'Pd(100) clean', 'Pd⁰', '80', '-434.380'),
    ('S2', '1 ML PdO(101)/Pd(100) (√5×√5)R27°', 'Pd⁰+Pd²⁺', '112', '-618.565'),
    ('S3', 'PdO(100) O-rich top (stoichiometric)', 'Pd²⁺', '128', '-724.103'),
    ('S3b', 'PdO(100) Pd-rich top', 'Pd²⁺', '104', '-570.772'),
    ('S4', 'PdO₂(110) stoichiometric', 'Pd⁴⁺', '144', '-788.493'),
]:
    row = table.add_row().cells
    for i, val in enumerate(r): row[i].text = val

# =========== 3. MLIP RANKING RESULTS ===========
doc.add_page_break()
add_h('3. MLIP Ranking — Phase 1+2+3 Results', 1)
add_para(
    "Phase 1 (단일 흡착 CO*, CH3O* 2,516 후보) + Phase 2 (co-ads SetA reactive 37,956) + "
    "Phase 3 (SetTS + SetB 6,314). 총 ~46,786 relaxation. Phase 2/3 에서 PBC fragmentation "
    "silent bug 발견 후 direct-distance refilter 적용 (~50% 생존). "
    "Phase 3 SetTS pool 은 가이드 T2.5 재해석으로 사실상 불필요 판정 (NEB 가 endpoint 에서 saddle 탐색)."
)
add_h('3.1 표면별 chemistry summary (MLIP top-1)', 2)
table = doc.add_table(rows=1, cols=7)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['Surface', '산화', 'CO d_min (Å)', 'CO E_range (meV)',
                       'CH₃O d_min (Å)', 'co-ads d_react (Å)', 'Case 예상']):
    hdr[i].text = h
for r in [
    ('S1', 'Pd⁰', '1.97 ✓', '1692', '2.11 ✓', '3.10 (SetA)', 'A'),
    ('S2', 'Pd⁰+Pd²⁺', '2.01 ✓', '818', '2.14 ✓', '5.26 (drift→B)', 'A/B'),
    ('S3', 'Pd²⁺ O-top', '2.46 ⚠', '126', '2.79 ⚠', '1.34 (product) ⭐', 'C'),
    ('S3b', 'Pd²⁺ Pd-top', '3.54 ❌', '2314', '2.55 ⚠', '5.30 (drift→B)', 'A/B'),
    ('S4', 'Pd⁴⁺', '4.05 ❌', '3184', '0.98 ❌', '1.33 (broken)', 'C/D'),
]:
    row = table.add_row().cells
    for i, val in enumerate(r): row[i].text = val

doc.add_picture(str(FIG/'F01_oxidation_trend.png'), width=Inches(6.0))
add_caption('Figure 1. Pd⁰ → Pd⁴⁺ 산화 진행에 따른 chemisorption 강도 변화. CO* (Pd–C) 가 1.97 Å → 4.05 Å 로 단조 증가하며 chemisorbed band (1.85-2.15 Å) 이탈. CH₃O* (Pd–O) 도 유사 경향, S4 PdO₂ 에서 0.98 Å 까지 단축은 분해 신호. Shi 2024 의 anodic 산화 → CO* 안정성 손실 가설 정확 재현.')

doc.add_picture(str(FIG/'F02_descriptor_preview.png'), width=Inches(5.5))
add_caption('Figure 2. 예비 descriptor map (MLIP top-1, slab E 차감). 5 surfaces 가 Case A-D 영역 분포. DFT 후 T1.19 에서 ΔG_CO* vs ΔG_CH3O*^MeOH(U) 로 최종 갱신 예정.')

doc.add_picture(str(FIG/'F03_d_min_distributions.png'), width=Inches(6.5))
add_caption('Figure 3. Phase 1 (단일 흡착) d_min 분포. (a) CO* Pd–C, (b) CH₃O* Pd–O. S1, S3b 가 chemisorbed band 안에 가장 많이 모여있고, S3, S4 는 physisorbed/borderline 영역 집중 — Pd 산화 영향.')

doc.add_picture(str(FIG/'F04_d_reactive_coads.png'), width=Inches(6.5))
add_caption('Figure 4. Phase 2 filtered SetA co-adsorption: 반응 거리 d(C_CO ↔ O_methoxy) vs MACE+D3 ΔE. SetA band [2.1, 4.0] Å 음영 표시. S3, S4 에서 1.3-1.5 Å 영역 (product collapse) signal 관찰.')

# =========== 4. v2 SHORTLIST (GUIDE-STRICT) ===========
doc.add_page_break()
add_h('4. T1.15 v2 DFT Shortlist — Guide-Strict (48 candidates)', 1)
add_para(
    "docs/DMC_Pd_workplan.md §P1-C 의 MLIP ranking 우선순위 5개 모두 strict 적용:"
)
add_para(
    "  ① coadsorption total energy 낮은 구조\n"
    "  ② C_CO–O_CH3O = 2.0–3.5 Å (Set A reactive band)\n"
    "  ③ CO·CH3O 서로 다른 기능성 site (co-adsorption)\n"
    "  ④ S2 에서는 interface pair 우선\n"
    "  ⑤ O-rich PdO/PdO₂ 에서 CO₂-like 무너진 구조 → side-path 보관"
)

add_h('4.1 후보 분포', 2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['Surface', 'CO*', 'CH₃O*', 'co-ads (SetA)', 'Total']):
    hdr[i].text = h
for r in [('S1', '3', '3', '3', '9'), ('S2', '5', '5', '5', '15'),
          ('S3', '3', '3', '3', '9'), ('S3b', '3', '3', '3', '9'),
          ('S4', '3', '3', '3', '9'), ('Total', '17', '17', '15', '51')]:
    row = table.add_row().cells
    for i, val in enumerate(r): row[i].text = val
    if r[0] == 'Total':
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs: run.bold = True

doc.add_picture(str(FIG/'F05_site_distribution.png'), width=Inches(6.5))
add_caption('Figure 5. v2 shortlist site type 분포. (a) CO* 의 C anchor 위치, (b) CH₃O* 의 methoxy O anchor 위치, (c) co-ads (CO·CH3O 두 anchor 합산). 5 surfaces 모두 다양한 site type 포함: atop_Pd, bridge_Pd-Pd, hollow_3Pd (S1, S3b), atop_O, bridge_Pd-O (S3, S4 oxide). 가이드 ③ "서로 다른 기능성 site" strict 적용 결과.')

# =========== 5. REPRESENTATIVE STRUCTURES ===========
doc.add_page_break()
add_h('5. Representative Structures (48 DFT candidates 시각화)', 1)
add_para(
    "각 surface 별 v2 shortlist candidates 의 top view (왼쪽) + side view (오른쪽). "
    "Surface 별로 페이지 구성:"
)

for sid, name in [('S1', 'Pd(100) [Pd⁰]'),
                  ('S2', 'PdO(101)/Pd(100) (√5×√5)R27° [Pd⁰+Pd²⁺ interface]'),
                  ('S3', 'PdO(100) O-rich top [Pd²⁺]'),
                  ('S3b', 'PdO(100) Pd-rich top [Pd²⁺]'),
                  ('S4', 'PdO₂(110) stoichiometric [Pd⁴⁺]')]:
    add_h(f'5.{["S1","S2","S3","S3b","S4"].index(sid)+1}  {sid} — {name}', 2)
    fig_path = FIG / f'F0{6 + ["S1","S2","S3","S3b","S4"].index(sid)}_representative_{sid}.png' if sid != 'S4' else FIG / 'F10_representative_S4.png'
    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(6.5))
        add_caption(f'Figure 6+. {sid} v2 shortlist 후보들. 좌: top view, 우: side view. 각 cell 좌상단 label = rank + site type.')
    doc.add_paragraph()

# =========== 6. COMMITTEE VALIDATION HISTORY ===========
doc.add_page_break()
add_h('6. Committee Validation History (Paimon-Extended 5-Judge)', 1)
add_para(
    "Paimon framework (Park et al., arXiv:2606.09422) 영감으로 5-judge blind parallel committee. "
    "judge-methods (룰 준수), judge-physics (chemistry sanity), judge-statistics (분포 validity), "
    "judge-silent-error (Paimon §2.2 — 'plausible but wrong' 감지), judge-malicious (Paimon §3.3 — "
    "trust audit). 모든 high-stakes 산출물 검증, silent bug 2회 발견 + fix."
)
doc.add_picture(str(FIG/'F11_committee_timeline.png'), width=Inches(6.5))
add_caption('Figure 11. 11개 committee verification cycle. P=Pass, P-c=Pass-with-caveats, C=Concern, R=Reject. Phase 2/3 raw 의 silent-error judge REJECT 가 결정적 PBC bug fix 트리거. T1.15 v2 site-strict 에서 모든 judge Pass.')

# =========== 7. VASP SETUP ===========
add_h('7. VASP Input Verification', 1)
add_para('INCAR 표준 (workplan §1):', bold=True, size=12)
add_para(
    "  ENCUT=520, PREC=Accurate, LASPH=.TRUE., ADDGRID=.TRUE., ISPIN=2\n"
    "  IVDW=12 (D3-BJ),  EDIFF=1e-06,  IBRION=2,  NSW=300,  ISIF=2 (ionic only)\n"
    "  EDIFFG=-0.03,  ISYM=0,  LDIPOL=.TRUE.,  IDIPOL=3,  KSPACING=0.25\n\n"
    "Material-specific:\n"
    "  S1 (Pd metal) : ISMEAR=1, SIGMA=0.10\n"
    "  S2-S4 (oxides): ISMEAR=0, SIGMA=0.05\n\n"
    "POTCAR: Pd_pv (28Jan2005, 16 valence) + O (08Apr2002) + C/H for co-ads.\n"
    "POTCAR library: /home/hyunjin/POTENTIAL/potpaw_PBE/"
)

# =========== 8. COST ===========
add_h('8. DFT 비용 견적', 1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['단계', '후보', '시간/job', 'GPU-hr 합계', 'Wall (2 GPUs)']):
    hdr[i].text = h
for r in [('T1.16 Level 1 (vacuum)', '48', '6-12 hr', '288-576', '6.0-12.0 day'),
          ('T1.17 Level 2 (VASPsol SP)', '48', '1-2 hr', '48-96', '1.0-2.0 day'),
          ('Total', '96 calc', '', '336-672', '7.0-14.0 day')]:
    row = table.add_row().cells
    for i, val in enumerate(r): row[i].text = val
    if r[0] == 'Total':
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs: run.bold = True

add_para(
    "GPU 자원: 2× RTX 6000 Ada Generation. SLURM debug partition. NVHPC compiler + MPI."
)

# =========== 9. VERSION MANAGEMENT ===========
doc.add_page_break()
add_h('9. Version Management — Pattern A (Gate-Based Branches)', 1)
add_para(
    "Advisor 이태훈 교수 권고 (2026-06-22): gate 별 long-running branch 영구 보존."
)
add_para('GitHub repositories (양쪽 mirror):', bold=True, size=12)
add_para(
    "  • https://github.com/hyunjinlee12/DMC      (개인 backup, origin)\n"
    "  • https://github.com/DDMM-dgist/Pd-DMC     (research group, ddmm)"
)

add_h('Branch 구조', 2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['Branch', 'Commit', '내용', '상태']):
    hdr[i].text = h
for r in [
    ('main', '4da551a', 'G2 통과 baseline (stable trunk)', 'stable'),
    ('g1-bulk-passed', '458ae51', 'G1 통과 marker (Pd, PdO, PdO₂ bulk)', 'frozen'),
    ('g2-slab-passed', '1d6eaa8', 'G2 통과 (= tag v0.2-G2-passed)', 'frozen'),
    ('g3-mlip-shortlist', '870b2a3', 'T1.14-T1.15 MLIP + DFT shortlist v2', 'active'),
    ('g3-dft (예정)', '—', 'T1.16-T1.20 DFT 결과', 'to create'),
    ('g4-neb (예정)', '—', 'workplan Phase 2 (NEB)', 'to create'),
]:
    row = table.add_row().cells
    for i, val in enumerate(r): row[i].text = val

add_para('Tags (immutable):', bold=True, size=12)
add_para(
    "  • v0.2-G2-passed (1d6eaa8) — G2 시점 영구 보존\n"
    "  • v0.3-G3-passed (예정) — G3 통과 시\n"
    "  • v1.0-final (예정) — G4 통과, 프로젝트 종료 시"
)

# =========== 10. NEXT STEPS ===========
add_h('10. 후속 단계 (T1.16 → G4)', 1)
add_para(
    "T1.16 종료 후 자동 sequence:\n\n"
    "  1. 48 OUTCAR 자동 검증 (free atom force < 0.03 eV/Å, 수렴 도달)\n"
    "  2. T1.17 Level 2 VASPsol single-point (LSOL=.TRUE., EB_K=32.6)\n"
    "  3. T1.18 adsorption energy table (ΔG_CO*, ΔG_CH3O*^MeOH(U) — CHE 보정)\n"
    "  4. T1.19 descriptor map (DFT 기반)\n"
    "  5. T1.20 Case A-D 분류 + Phase 2 surfaces 선정\n"
    "  6. G3 게이트 통과 → checkpoint C/D/E\n"
    "  7. g3-dft branch → main merge + tag v0.3-G3-passed\n"
    "  8. g4-neb 새 branch → workplan Phase 2 (CI-NEB TS1/TS2 + side-path)\n"
    "  9. T2.5-T2.8 → G4 통과 → tag v1.0-final → 프로젝트 종료"
)

# =========== 11. DECISION REQUEST ===========
add_h('11. Decision Request', 1)
add_para('Advisor 검토 요청 사항:', bold=True, size=12)
add_para(
    "  1. T1.16 Level 1 vacuum DFT 진입 승인 (48 jobs, ~6-12 day wall time on 2 GPUs)\n"
    "  2. v2 shortlist 가이드-strict pick 적용 OK 인지 (48 candidates, 표면별 분포 표 참조)\n"
    "  3. Pattern A version management 의도 일치 확인 (gate 별 long-running branch + DDMM mirror)\n"
    "  4. S4 PdO₂(110) literature gap 명시 동의 — 'representative Pd⁴⁺ surface' 로 해석 가능 여부"
)

# =========== 12. REFERENCES ===========
add_h('12. References', 1)
refs = [
    "Shi et al. \"Stabilization of Pd⁰ by Cu Alloying: Pd₃Cu Electrocatalyst for Anodic Methanol Carbonylation\", Angew. Chem. Int. Ed. 2024, doi:10.1002/anie.202401311",
    "Park, Chung, You, Kim, Ju, Han. \"A Robust Agentic Framework for Expert-Level Automation of Atomistic Simulations (Paimon)\", arXiv:2606.09422, 2026",
    "Lundgren et al. \"The Pd(100)-(√5×√5)R27°-O surface oxide revisited\", arXiv:cond-mat/0304107 (2003)",
    "Rogal, Reuter, Scheffler. \"Thermodynamic stability of PdO surfaces\", PRB 2004, arXiv:cond-mat/0310235",
    "Reuter et al. \"CO oxidation at Pd(100): First-principles constrained thermodynamics\", PRL 98, 046101 (2007)",
    "MACE-MP foundation model — Batatia et al., arXiv:2401.00096; https://mace-docs.readthedocs.io",
    "AutoAdsorbate — ACS Catal. 2025, doi:10.1021/acscatal.5c06553",
    "cuEquivariance (NVIDIA) — https://github.com/NVIDIA/cuEquivariance",
]
for r in refs:
    p = doc.add_paragraph(r, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Arial'; run.font.size = Pt(10)

doc.save(str(OUT))
print(f'✓ Report: {OUT}')
print(f'  Size: {OUT.stat().st_size/1024:.0f} KB')
