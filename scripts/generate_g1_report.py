"""Generate G1 Report as Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

PROJECT = Path(__file__).resolve().parent.parent
OUT = PROJECT / "calculations" / "G1_bulk"


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
        if highlight_rows and r_idx in highlight_rows:
            for c_idx in range(len(row)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "DAEEF3")
    return table


def add_figure(doc, path, width, caption):
    if not Path(path).exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("G1 — Bulk Optimization Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Pd/PdO/PdO₂ DMC DFT Study  |  2026-06-04")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # 1. Objective
    doc.add_heading("1. Objective", level=1)
    doc.add_paragraph(
        "Pd, PdO, PdO₂ bulk 구조를 PBE+D3/520 eV 수준에서 "
        "최적화하고, ENCUT\xb7k-mesh 수렴 검증 + "
        "실험/문헌 격자상수 비교를 통해 "
        "slab 제작의 기준 구조를 확보한다."
    )

    # 2. Results
    doc.add_heading("2. Results", level=1)

    # 2.1 Initial structures
    doc.add_heading("2.1 Initial Structures", level=2)
    doc.add_paragraph(
        "Materials Project (MP) API로 초기 구조를 확보하고, "
        "conventional standard cell로 변환 후 사용."
    )
    add_table(doc,
        ["Material", "MP-ID", "Space Group", "ICSD", "Source"],
        [
            ["Pd", "mp-2", "Fm-3m (fcc)", "—", "MP (exp. XRD)"],
            ["PdO", "mp-1336", "P4₂/mmc", "—", "MP (exp. XRD)"],
            ["PdO₂", "mp-1018886", "P4₂/mnm (rutile)", "647283", "MP ← ICSD [3]"],
        ])

    # 2.2 INCAR
    doc.add_heading("2.2 Key INCAR Parameters", level=2)
    doc.add_paragraph(
        "금속(Pd)과 산화물(PdO, PdO₂)의 주요 INCAR 차이:"
    )
    add_table(doc,
        ["Tag", "Pd (금속)", "PdO · PdO₂ (산화물)", "비고"],
        [
            ["ISMEAR", "1", "0", "금속 Methfessel-Paxton / 산화물 Gaussian"],
            ["SIGMA", "0.10", "0.05", "산화물 band gap → 좁은 smearing"],
            ["IVDW", "12", "12", "D3-BJ (공통)"],
            ["ISIF", "3", "3", "full cell + ionic relax (bulk)"],
            ["ENCUT", "520 eV", "520 eV", "수렴 테스트 확정 (공통)"],
        ])
    doc.add_paragraph()

    # 2.3 Convergence
    doc.add_heading("2.3 Convergence Test", level=2)
    add_table(doc,
        ["ENCUT (eV)", "Pd (meV/atom)", "PdO (meV/atom)", "PdO₂ (meV/atom)"],
        [
            ["400", "+1.97", "−5.32", "−8.25"],
            ["450", "+0.68", "−0.32", "−0.42"],
            ["500", "+0.26", "+0.77", "+1.25"],
            ["520", "+0.17", "+0.72", "+1.16"],
            ["550", "+0.23", "+0.63", "+0.94"],
            ["600", "ref", "ref", "ref"],
        ],
        highlight_rows=[3])
    doc.add_paragraph()
    doc.add_paragraph("Adopted k-mesh:  Pd 12\xd712\xd712  \xb7  PdO 8\xd78\xd76  \xb7  PdO₂ 6\xd76\xd78")

    add_figure(doc,
        OUT / "convergence" / "results" / "convergence_plots.png", 6.0,
        "Figure 1. ENCUT (top) and k-mesh (bottom) convergence. Dashed line = \xb11 meV/atom.")

    # 2.4 Optimized structures
    doc.add_heading("2.4 Optimized Lattice Parameters", level=2)
    add_table(doc,
        ["Material", "MP-ID", "a_DFT (\xc5)", "c_DFT (\xc5)",
         "a_ref (\xc5)", "c_ref (\xc5)", "err_a", "err_c"],
        [
            ["Pd", "mp-2", "3.8907", "—",
             "3.890 [1]", "—", "+0.02%", "—"],
            ["PdO", "mp-1336", "3.0536", "5.4058",
             "3.043 [2]", "5.328 [2]", "+0.35%", "+1.46%"],
            ["PdO₂", "mp-1018886", "4.5424", "3.1772",
             "4.486 [3]*", "3.103 [3]*", "+1.26%", "+2.39%"],
        ])
    doc.add_paragraph()
    p = doc.add_paragraph(
        "* PdO₂: Shaplygin et al. (1978)이 고압 합성하여 "
        "ICSD 647283으로 등록했으나, "
        "원문(러시아어)의 실험 격자상수가 "
        "영문 문헌에서 직접 인용되지 않음. "
        "MP mp-1018886 (GGA-PBE, ICSD 647283 기반)을 "
        "computational reference로 사용. "
        "Matar et al. (2011)도 GGA+U로 연구하여 "
        "실험과 유사한 격자상수를 보고."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True

    add_figure(doc,
        OUT / "G1_lattice_comparison.png", 5.0,
        "Figure 2. DFT vs reference lattice parameters.")

    # 3. Conclusion
    doc.add_heading("3. Conclusion", level=1)
    p = doc.add_paragraph()
    run = p.add_run("G1 PASS.")
    run.bold = True
    p.add_run(
        " 3종 bulk 모두 수렴 확인, "
        "격자 오차 ≤ 2.4%. Slab 제작 (G2) 진행 가능."
    )

    # References
    doc.add_heading("References", level=2)
    refs = [
        "[1] Kittel, C. Introduction to Solid State Physics, 8th ed. (Wiley). "
        "Pd fcc a = 3.89 \xc5.",

        "[2] Waser, J., Levy, H. A., Peterson, S. W. "
        "“The structure of PdO.” "
        "Acta Cryst. 1953, 6, 661–663. "
        "DOI: 10.1107/S0365110X53001800",

        "[3] Shaplygin, I. S., Aparnikov, G. L., Lazarev, V. B. "
        "“Formation of Palladium Dioxide at High Pressures.” "
        "Zh. Neorg. Khim. 1978, 23(4), 884–886. "
        "ICSD 647283. MP mp-1018886 (GGA-PBE relaxation of ICSD structure).",

        "[4] Matar, S. F., Demazeau, G., M\xf6ller, M. H., P\xf6ttgen, R. "
        "“Electronic structure and equation of state of PdO₂ from ab initio.” "
        "Chem. Phys. Lett. 2011, 508, 215–218. "
        "DOI: 10.1016/j.cplett.2011.04.054",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.runs[0].font.size = Pt(9)

    out_path = OUT / "G1_Report.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
