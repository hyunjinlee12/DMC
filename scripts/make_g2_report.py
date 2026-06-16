"""Generate G2 Slab Construction Report (Word) with side/perspective views.

Usage:
    conda run -n pddmc python scripts/make_g2_report.py
"""

import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from ase.io import read
from ase.constraints import FixAtoms

PROJECT = Path(__file__).resolve().parent.parent
G2 = PROJECT / "calculations" / "G2_slab"

SLABS = [
    {
        "id": "S1",
        "dir": "S1_Pd100",
        "name": "Pd(100)",
        "ox": "Pd⁰",
        "supercell": "p(4×4), 5 layers",
        "bulk_source": "Pd fcc (a = 3.891 Å, G1 PBE+D3)",
        "method": "ASE fcc100 builder → p(4×4) 5-layer supercell",
        "rationale": (
            "Pd nanocube {100} facet의 metallic baseline. "
            "산화 전 순수 금속 Pd 표면으로, Shi 2024 (Angew. Chem.)의 "
            "pure Pd 계산과 직접 비교 가능. DMC carbonylation 반응의 "
            "기준점(reference)으로 사용."
        ),
        "fix": "하부 2 layers 고정 (substrate 역할)",
        "notes": "가장 단순한 구조. fcc(100) 정방형 배열.",
    },
    {
        "id": "S2",
        "dir": "S2_PdO101_Pd100",
        "name": "PdO(101)/Pd(100)",
        "ox": "Pd⁰/Pd²⁺",
        "supercell": "(√5×√5)R27° × 2×2",
        "bulk_source": "Pd fcc (a = 3.891 Å) + PdO (a = 3.054, c = 5.406 Å)",
        "method": (
            "Pd(100) 4-layer substrate에 vacuum 부여 후 (√5×√5)R27° × 2×2 supercell 변환 "
            "(P = [[4,2,0],[-2,4,0],[0,0,1]], det=20). "
            "PdO(101) 단일층을 crystallographic 좌표로 배치. "
            "Todorova et al. (Surf. Sci. 541, 101, 2003) 모델 기반. "
            "※ vacuum 없이 supercell 시 top layer가 z=0으로 wrapping되는 버그 수정 완료."
        ),
        "rationale": (
            "Pd(100) 산화 초기 단계에서 형성되는 (√5×√5)R27° surface oxide의 "
            "표준 모델. LEED + DFT + STM으로 확립된 구조. "
            "Pd/PdO interface에서 CO*와 CH₃O*를 동시에 안정화하는 "
            "bifunctional site 가능성 검증 (Case B). "
            "Strain: PdO(101) → √5 cell 매칭 시 a +0.7%, b −0.4%."
        ),
        "fix": "하부 2 layers 고정",
        "notes": (
            "Oxide monolayer는 근사 좌표 — DFT relaxation 필수. "
            "VESTA 검증 후 제출."
        ),
    },
    {
        "id": "S3",
        "dir": "S3_PdO100",
        "name": "PdO(100)",
        "ox": "Pd²⁺",
        "supercell": "2×2 supercell, O-rich termination",
        "bulk_source": "PdO tetragonal P4₂/mmc (a = 3.054, c = 5.406 Å, G1 PBE+D3)",
        "method": (
            "pymatgen SlabGenerator로 (100) 방향 절단 "
            "(min_slab_size=10 Å, min_vacuum=20 Å). "
            "P4₂/mmc 대칭 상 유일한 termination (O-rich top) → 2×2 supercell."
        ),
        "rationale": (
            "완전 산화된 bulk PdO의 thermodynamic/O-rich limit. "
            "CO adsorption이 약해질 수 있어 DMC-inactive 가능성 검증 (Case C). "
            "P4₂/mmc 대칭 제약으로 (100) termination은 1종류만 존재."
        ),
        "fix": "하부 ~35% 원자 고정 (z_fraction=0.35)",
        "notes": "Termination 유일성 확인 완료 (min_slab_size 8–20 Å 테스트).",
    },
    {
        "id": "S4",
        "dir": "S4_PdO2_110",
        "name": "PdO₂(110)",
        "ox": "Pd⁴⁺",
        "supercell": "3×2 supercell, stoichiometric (O/Pd = 2.0)",
        "bulk_source": "PdO₂ rutile P4₂/mnm (a = 4.542, c = 3.177 Å, G1 PBE+D3)",
        "method": (
            "pymatgen SlabGenerator로 (110) 방향 절단. "
            "2개 termination 중 stoichiometric (O/Pd=2.0) + Pd-exposed 선택. "
            "In-plane > 8 Å 확보를 위해 3×2 supercell."
        ),
        "rationale": (
            "Pd⁴⁺ high-valence oxide reference. "
            "극단적 산화 조건에서의 DMC 반응성 검증 (Case D). "
            "Rutile PdO₂는 고압 합성 상이나, anodic 조건에서 "
            "PdO₂-like overlayer 형성 가능성 있음."
        ),
        "fix": "하부 ~35% 원자 고정 (z_fraction=0.35)",
        "notes": (
            "2개 stoichiometric termination 존재: "
            "Term 0 (Pd-exposed, 채택) vs Term 1 (O-exposed, 대안)."
        ),
    },
]


def get_slab_info(poscar_path):
    atoms = read(str(poscar_path))
    syms = atoms.get_chemical_symbols()
    n_pd = syms.count("Pd")
    n_o = syms.count("O")
    a = np.linalg.norm(atoms.cell[0])
    b = np.linalg.norm(atoms.cell[1])
    c = np.linalg.norm(atoms.cell[2])
    z = atoms.positions[:, 2]
    thick = z.max() - z.min()
    vacuum = c - thick

    n_fixed = 0
    for con in atoms.constraints:
        if isinstance(con, FixAtoms):
            n_fixed += len(con.index)

    z_sorted = np.sort(z)
    layers = [[z_sorted[0]]]
    for zz in z_sorted[1:]:
        if zz - np.mean(layers[-1]) > 0.5:
            layers.append([zz])
        else:
            layers[-1].append(zz)

    return {
        "n_atoms": len(atoms),
        "n_pd": n_pd,
        "n_o": n_o,
        "a": a, "b": b, "c": c,
        "thick": thick,
        "vacuum": vacuum,
        "n_fixed": n_fixed,
        "n_layers": len(layers),
        "fix_pct": n_fixed / len(atoms) * 100,
    }


def set_cell_shading(cell, color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title = doc.add_heading("G2 Clean Slab Construction Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pd/PdO/PdO₂ DMC Formation DFT Study — Surface Model Summary")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("전북대학교 | 이태훈 연구실 | 2026-06-05")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 1. Overview
    doc.add_heading("1. 표면 모델 개요", level=1)

    doc.add_paragraph(
        "본 프로젝트는 Pd nanocube {100} facet이 anodic methanol carbonylation 조건에서 "
        "산화·재구성될 때 DMC 생성 경로가 어떻게 변화하는지를 DFT로 검증한다. "
        "이를 위해 산화 진행 축(oxidation hierarchy)을 따라 4개 표면 모델을 설정하였다."
    )

    # Hierarchy diagram
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "S1 (Pd⁰)  →  S2 (Pd⁰/Pd²⁺)  →  S3 (Pd²⁺)  →  S4 (Pd⁴⁺)\n"
        "metallic      surface oxide      bulk oxide     high-valence"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_paragraph(
        "Phase 1 종료 시 각 표면의 CO*·CH₃O* 흡착 에너지를 비교하여 "
        "Case A–D 프레임으로 DMC 활성 표면을 판정하고, Phase 2에서 "
        "유리한 표면만 TS barrier를 계산한다."
    )

    # Case table
    doc.add_heading("1.1 해석 프레임 (Case A–D)", level=2)
    t = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Case", "시나리오", "의미"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            p.runs[0].bold = True
    cases = [
        ("A", "Pd(100) 최유리", "순수 금속 Pd에서 DMC 생성 최적"),
        ("B", "PdO(101)/Pd(100) 최유리", "부분 산화 bifunctional interface가 DMC 촉진"),
        ("C", "bulk PdO(100) DMC-inactive", "완전 산화 시 CO* 약화 → side-path 우세"),
        ("D", "PdO₂(110) DMC-inactive", "고산화 극한 → DMC 비활성"),
    ]
    for i, (c, s, m) in enumerate(cases):
        t.rows[i + 1].cells[0].text = c
        t.rows[i + 1].cells[1].text = s
        t.rows[i + 1].cells[2].text = m

    # 2. Summary table
    doc.add_heading("2. Slab 구조 요약", level=1)

    t = doc.add_table(rows=5, cols=8, style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Surface", "산화 상태", "Atoms", "Pd", "O",
               "Cell (a×b Å)", "Thickness (Å)", "Fixed"]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)

    for row_idx, slab in enumerate(SLABS):
        poscar = G2 / slab["dir"] / "POSCAR"
        info = get_slab_info(poscar)
        vals = [
            f"{slab['id']} {slab['name']}",
            slab["ox"],
            str(info["n_atoms"]),
            str(info["n_pd"]),
            str(info["n_o"]),
            f"{info['a']:.1f} × {info['b']:.1f}",
            f"{info['thick']:.1f}",
            f"{info['n_fixed']}/{info['n_atoms']} ({info['fix_pct']:.0f}%)",
        ]
        for col_idx, val in enumerate(vals):
            cell = t.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # Computational parameters
    doc.add_heading("2.1 공통 계산 파라미터", level=2)
    params = [
        ("ENCUT", "520 eV"),
        ("KSPACING", "0.25 Å⁻¹ (전체 통일)"),
        ("XC Functional", "PBE + D3-BJ (Grimme)"),
        ("POTCAR", "Pd_pv (16e⁻) + O (standard)"),
        ("Vacuum", "20 Å (asymmetric, top side only)"),
        ("Dipole correction", "LDIPOL=.TRUE., IDIPOL=3"),
        ("ISIF", "2 (ionic relaxation only)"),
        ("ISYM", "0 (no symmetry)"),
    ]
    t = doc.add_table(rows=len(params) + 1, cols=2, style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.rows[0].cells[0].text = "Parameter"
    t.rows[0].cells[1].text = "Value"
    for p in t.rows[0].cells[0].paragraphs:
        p.runs[0].bold = True
    for p in t.rows[0].cells[1].paragraphs:
        p.runs[0].bold = True
    for i, (k, v) in enumerate(params):
        t.rows[i + 1].cells[0].text = k
        t.rows[i + 1].cells[1].text = v

    # 3. Per-surface details with images
    doc.add_heading("3. 표면별 상세 설명 및 구조", level=1)

    for slab in SLABS:
        doc.add_heading(f"{slab['id']}: {slab['name']} ({slab['ox']})", level=2)

        poscar = G2 / slab["dir"] / "POSCAR"
        info = get_slab_info(poscar)

        # Rationale
        doc.add_heading("선택 근거", level=3)
        doc.add_paragraph(slab["rationale"])

        # Construction
        doc.add_heading("제작 방법", level=3)
        doc.add_paragraph(f"Bulk source: {slab['bulk_source']}")
        doc.add_paragraph(f"Supercell: {slab['supercell']}")
        doc.add_paragraph(f"Method: {slab['method']}")
        doc.add_paragraph(f"Constraint: {slab['fix']}")

        # Structural info
        doc.add_heading("구조 정보", level=3)
        t = doc.add_table(rows=2, cols=6, style="Light Grid Accent 1")
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Atoms", "Layers", "a (Å)", "b (Å)",
                                "Thickness (Å)", "Vacuum (Å)"]):
            t.rows[0].cells[i].text = h
            for p in t.rows[0].cells[i].paragraphs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)
        for i, v in enumerate([
            str(info["n_atoms"]),
            str(info["n_layers"]),
            f"{info['a']:.2f}",
            f"{info['b']:.2f}",
            f"{info['thick']:.2f}",
            f"{info['vacuum']:.1f}",
        ]):
            t.rows[1].cells[i].text = v
            for p in t.rows[1].cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

        # Images: side + perspective
        doc.add_heading("구조 시각화", level=3)

        img_table = doc.add_table(rows=2, cols=2)
        img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        side_img = G2 / slab["dir"] / f"{slab['dir']}_side.png"
        persp_img = G2 / slab["dir"] / f"{slab['dir']}_persp.png"

        img_table.rows[0].cells[0].text = "Side View"
        img_table.rows[0].cells[1].text = "Perspective View"
        for c in img_table.rows[0].cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        if side_img.exists():
            p = img_table.rows[1].cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(side_img), width=Inches(2.8))

        if persp_img.exists():
            p = img_table.rows[1].cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(persp_img), width=Inches(2.8))

        # Notes
        if slab["notes"]:
            doc.add_heading("비고", level=3)
            doc.add_paragraph(slab["notes"])

        doc.add_page_break()

    # 4. Structural verification
    doc.add_heading("4. 구조 검증 결과", level=1)

    doc.add_paragraph(
        "각 슬랩의 레이어 구조, 원자 분포, 층간 거리를 검증하였다. "
        "S2는 초기 빌드 과정에서 make_supercell의 periodic wrapping으로 "
        "기판 바닥 2개 레이어가 동일 z에 붕괴되는 버그가 발견되어 수정하였다."
    )

    for slab in SLABS:
        poscar = G2 / slab["dir"] / "POSCAR"
        atoms = read(str(poscar))
        syms_v = atoms.get_chemical_symbols()
        z_v = atoms.positions[:, 2]
        info = get_slab_info(poscar)

        doc.add_heading(f"{slab['id']} {slab['name']}", level=3)

        z_sorted_v = np.sort(z_v)
        layers_v = [[z_sorted_v[0]]]
        for zz in z_sorted_v[1:]:
            if zz - np.mean(layers_v[-1]) > 0.5:
                layers_v.append([zz])
            else:
                layers_v[-1].append(zz)

        n_fixed_v = 0
        fix_z_max_v = 0
        for con in atoms.constraints:
            if isinstance(con, FixAtoms):
                n_fixed_v = len(con.index)
                fix_z_max_v = max(z_v[i] for i in con.index)

        n_rows = len(layers_v) + 1
        t = doc.add_table(rows=n_rows, cols=5, style="Light Grid Accent 1")
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Layer", "z (Å)", "Atoms", "Composition", "Status"]):
            t.rows[0].cells[ci].text = h
            for p in t.rows[0].cells[ci].paragraphs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(8.5)

        for li, layer in enumerate(layers_v):
            z_avg = np.mean(layer)
            layer_idx = [j for j in range(len(atoms)) if abs(z_v[j] - z_avg) < 0.3]
            pd_n = sum(1 for j in layer_idx if syms_v[j] == "Pd")
            o_n = sum(1 for j in layer_idx if syms_v[j] == "O")
            comp_parts = []
            if pd_n:
                comp_parts.append(f"Pd={pd_n}")
            if o_n:
                comp_parts.append(f"O={o_n}")
            is_fixed = all(z_v[j] <= fix_z_max_v + 0.1 for j in layer_idx) and n_fixed_v > 0

            vals = [
                f"L{li+1}",
                f"{z_avg:.2f}",
                str(len(layer)),
                ", ".join(comp_parts),
                "Fixed" if is_fixed else "Relaxed",
            ]
            for ci, v in enumerate(vals):
                t.rows[li + 1].cells[ci].text = v
                for p in t.rows[li + 1].cells[ci].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)

        # Interlayer spacings
        centers_v = sorted(np.mean(L) for L in layers_v)
        spacings = [f"L{i+1}→L{i+2}: {centers_v[i+1]-centers_v[i]:.3f} Å"
                    for i in range(len(centers_v) - 1)]
        doc.add_paragraph("층간 거리: " + " | ".join(spacings))

    # 5. Oxidation hierarchy summary
    doc.add_heading("5. 산화 진행 축 요약", level=1)

    doc.add_paragraph(
        "아래 표는 4개 표면의 산화 상태 진행을 정리한 것이다. "
        "Phase 1에서 각 표면의 CO*·CH₃O* 흡착 에너지를 비교한 후, "
        "Case A–D 프레임으로 DMC 활성 표면을 판정한다."
    )

    t = doc.add_table(rows=5, cols=5, style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Surface", "산화 상태", "역할", "Case", "DMC 전망"]):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            p.runs[0].bold = True

    rows_data = [
        ("S1 Pd(100)", "Pd⁰", "Metallic baseline", "A", "기준점 — pure Pd 비교"),
        ("S2 PdO(101)/Pd(100)", "Pd⁰/Pd²⁺", "부분 산화 bifunctional", "B",
         "CO*+CH₃O* 동시 안정화 가능"),
        ("S3 PdO(100)", "Pd²⁺", "Bulk oxide / O-rich", "C",
         "CO* 약화 → inactive 가능"),
        ("S4 PdO₂(110)", "Pd⁴⁺", "High-valence reference", "D",
         "극한 산화 → inactive 가능"),
    ]
    for i, vals in enumerate(rows_data):
        for j, v in enumerate(vals):
            t.rows[i + 1].cells[j].text = v

    # 6. References
    doc.add_heading("6. 주요 참고문헌", level=1)
    refs = [
        "Shi et al., Angew. Chem. Int. Ed., 2024 — Pd₃Cu DMC electrocatalyst, "
        "TS barrier benchmark (TS1=1.08, TS2=0.85 eV for pure Pd).",
        "Todorova et al., Surf. Sci. 541, 101, 2003 (cond-mat/0304107) — "
        "Pd(100)-(√5×√5)R27°-O surface oxide structure.",
        "Rogal et al., Phys. Rev. B — Pd/PdO redox kinetic phase diagram, "
        "metal↔oxide dynamic switching under reaction conditions.",
        "Shaplygin et al., Zh. Neorg. Khim., 1978 (ICSD 647283) — "
        "PdO₂ rutile experimental lattice parameters.",
    ]
    for i, ref in enumerate(refs):
        doc.add_paragraph(f"[{i+1}] {ref}")

    # 7. Next steps
    doc.add_heading("7. 다음 단계", level=1)
    steps = [
        "VESTA로 4개 POSCAR 구조 육안 검증 (rumpling, Pd–O bond, coordination)",
        "DFT slab relaxation 제출 (VASP PBE+D3, ISIF=2)",
        "Gate G2: clean slab 물리량 검증 통과 후 adsorption sampling 착수",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Bullet")

    # Save
    out = G2 / "G2_Slab_Report_v3.docx"
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
